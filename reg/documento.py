from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from reg.listas import convenio_opciones
from babel.dates import format_date

def crear(paciente, presupuesto, cotizaciones):
    buff = BytesIO()
    document = Document()
    heading = document.add_heading('Clínica Dental Dr. Mauricio Martínez V.')
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheading = document.add_heading('Barros 357, Melipilla    Tel: +56 2 28313380', level=5)
    subheading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo = document.add_paragraph()
    titulo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.add_run(f'Presupuesto para {paciente.nombre} {paciente.apellido_paterno} {paciente.apellido_materno}').bold = True
    hoy = format_date(presupuesto.fecha_creacion, format="long", locale="es_CL")
    datos = document.add_paragraph(f'{hoy}, n° {presupuesto.id}')
    datos.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1 = document.add_paragraph()
    p1.add_run(f'Convenio: {dict(convenio_opciones)[presupuesto.convenio]}')
    #Tabla
    tabla = document.add_table(rows=1, cols=3)
    tabla.style = 'Medium List 2'
    hdr = tabla.rows[0].cells
    hdr[0].text = "Prestación"
    hdr[1].text = "Pieza"
    hdr[2].text = "Monto"
    for cotizado in cotizaciones:
        row = tabla.add_row().cells
        row[0].text = cotizado.prestacion.nombre
        row[1].text = cotizado.pieza
        row[2].text = str(cotizado.monto)
    monto_total = sum(x.monto for x in cotizaciones)
    rowf = tabla.add_row().cells
    rowf[1].text = "Monto total:"
    rowf[2].text = str(monto_total)
    document.add_page_break()
    document.save(buff)
    pdf = buff.getvalue()
    buff.close()
    return pdf
